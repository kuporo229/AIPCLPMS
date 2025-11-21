# app/utils.py

import os
import io
import json
from datetime import datetime
from threading import Thread
import google.generativeai as genai
from google.generativeai import GenerativeModel
from docx import Document
from flask import current_app, session, flash, url_for, redirect, jsonify, Response
import logging
import jwt
import time
from flask import current_app
import logging # Make sure logging is imported if not already
from werkzeug.utils import secure_filename
# Import the exception class from the library
from supabase import PostgrestAPIError
# Import the supabase client and other global variables from your app's main factory file
from app import supabase, STORAGE_BUCKET_NAME, ALLOWED_EXTENSIONS
import sys
from postgrest.exceptions import APIError as PostgrestAPIError

# Re-import static data from the app factory
from app import PROGRAM_OUTCOMES, COURSE_OUTCOMES, INSTITUTIONAL_OUTCOMES_HEADERS, PROGRAM_OUTCOMES_HEADERS

from supabase import create_client

# Normal client (anon/public key) - for user-facing queries
from app import supabase  
# Ensure logging is configured properly
if current_app:
    current_app.logger.setLevel(logging.INFO)
    if not current_app.logger.handlers:
        handler = logging.StreamHandler()
        handler.setLevel(logging.INFO)
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        current_app.logger.addHandler(handler)

# Create a separate service client (bypasses RLS)
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")  # store this in .env, NEVER frontend

supabase_service = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)


# app/utils.py

# ... existing imports (make sure you have `docx` imported)



# ... existing utility functions ...

# --- NEW DOCX UTILITY FUNCTIONS ---

def get_template_filepath():
    """Returns the absolute path to the official CLP template file."""
    # NOTE: You must ensure this file path is correct relative to where your Flask app runs
    # Assuming the template file is in the root directory for now
    template_filename = 'PBSIT-001-LP-20242 copy.docx'
    # Use current_app.root_path to find the file
    return os.path.join(os.path.dirname(__file__), 'clp_templates', 'PBSIT-001-LP-20242 copy.docx')



import base64

def load_template_content(filename):
    """
    Load the content of a .docx template as HTML/text for inline editing.
    Right now this just returns plain text, but you can later improve it
    with a docx → HTML converter (like python-docx, pydocx, or Mammoth).
    """
    template_path = os.path.join(os.path.dirname(__file__), 'clp_templates', filename)

    if not os.path.exists(template_path):
        return ""

    try:
        from docx import Document
        doc = Document(template_path)
        # Join paragraphs with line breaks
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        # Fallback: just return empty if parsing fails
        return ""
    
def get_template_content():
    """
    Reads DOCX and returns full HTML with headers, footers, text formatting, tables, and images.
    Works on all python-docx versions (no xpath).
    """
    try:
        doc_path = get_template_filepath()
        document = Document(doc_path)
        content = []

        # --- Headers ---
        for section in document.sections:
            header = section.header
            if header and header.paragraphs:
                content.append('<div style="border-bottom:1px solid #000; margin-bottom:10px;">')
                for para in header.paragraphs:
                    content.append(_format_paragraph_with_images(para))
                content.append('</div>')

        # --- Body paragraphs ---
        for para in document.paragraphs:
            content.append(_format_paragraph_with_images(para))

        # --- Tables ---
        for table in document.tables:
            table_html = '<table style="border-collapse:collapse; width:100%; margin:10px 0;">'
            for row in table.rows:
                table_html += "<tr>"
                for cell in row.cells:
                    cell_text = ''.join(_format_paragraph_with_images(p) for p in cell.paragraphs)
                    table_html += f'<td style="border:1px solid #444; padding:5px; vertical-align:top;">{cell_text}</td>'
                table_html += "</tr>"
            table_html += "</table>"
            content.append(table_html)

        # --- Footers ---
        for section in document.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                content.append('<div style="border-top:1px solid #000; margin-top:10px;">')
                for para in footer.paragraphs:
                    content.append(_format_paragraph_with_images(para))
                content.append('</div>')

        return "\n".join(content)

    except FileNotFoundError:
        current_app.logger.error(f"Template file not found at: {doc_path}")
        return "ERROR: Template file not found."
    except Exception as e:
        current_app.logger.error(f"Error reading DOCX: {e}")
        return f"ERROR: Could not read template content. Details: {e}"

# Ensure logging is configured (add if not present)
logger = logging.getLogger(__name__) # Or use current_app.logger if in request context

# ONLYOFFICE JWT Secret - Fetched from config
# ONLYOFFICE_JWT_SECRET = os.getenv("ONLYOFFICE_JWT_SECRET", "") # No longer needed here, use current_app.config

def generate_jwt_token(payload):
    """Generate JWT token for ONLYOFFICE if JWT secret is configured"""
    onlyoffice_jwt_secret = current_app.config.get("ONLYOFFICE_JWT_SECRET", "")
    if not onlyoffice_jwt_secret:
        # current_app might not be available if called outside request context
        # Use basic logging or handle appropriately
        logger.warning("ONLYOFFICE_JWT_SECRET not set in app config - cannot generate token.")
        return None

    try:
        # Add expiration time (e.g., 1 hour)
        payload['exp'] = int(time.time()) + 3600
        token = jwt.encode(payload, onlyoffice_jwt_secret, algorithm='HS256')
        # Use logger instead of current_app.logger if outside request context
        logger.info("Generated JWT token for ONLYOFFICE payload.")
        return token
    except Exception as e:
        # Use logger instead of current_app.logger
        logger.error(f"Error generating JWT: {e}")
        return None
    
def _format_paragraph_with_images(para):
    """
    Converts a DOCX paragraph into HTML, scanning for images by walking the XML tree
    (avoids xpath completely).
    """
    formatted_text = ""

    for run in para.runs:
        # --- Find embedded images by iterating XML elements ---
        for child in run._element.iter():
            if child.tag.endswith("blip"):  # <a:blip> element contains the image relationship ID
                rId = child.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if rId:
                    image_part = run.part.related_parts.get(rId)
                    if image_part:
                        image_data = image_part.blob
                        encoded_image = base64.b64encode(image_data).decode("utf-8")
                        formatted_text += f'<img src="data:image/png;base64,{encoded_image}" style="max-height:100px; margin:5px;">'

        # --- Handle text formatting ---
        if not run.text.strip():
            continue
        run_text = run.text
        if run.bold:
            run_text = f"<b>{run_text}</b>"
        if run.italic:
            run_text = f"<i>{run_text}</i>"
        if run.underline:
            run_text = f"<u>{run_text}</u>"
        formatted_text += run_text

    # --- Paragraph alignment ---
    align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
    align = para.alignment
    align_style = ""
    if align is not None and align in align_map:
        align_style = f"text-align:{align_map[align]};"

    return f'<p style="{align_style} margin:5px 0;">{formatted_text}</p>'

def overwrite_template_content(new_content):
    """Overwrites the content of the DOCX template with new text."""
    try:
        if not new_content or len(new_content.strip()) < 10:
            current_app.logger.warning("Refused to overwrite template: empty content received.")
            return False
        
        doc_path = get_template_filepath()
        
        # 1. Create a brand new, empty document
        new_document = Document()
        
        # 2. Split the content by double-newline to represent new paragraphs
        paragraphs = new_content.split('\n\n')
        
        # 3. Add each part as a new paragraph to the new document
        for text in paragraphs:
            # Replace single newlines within a paragraph with a line break (for soft wrapping)
            text_with_breaks = text.replace('\n', '\n') 
            new_document.add_paragraph(text_with_breaks)
            
        # 4. Save the new document, overwriting the old template
        new_document.save(doc_path)
        return True
    except Exception as e:
        current_app.logger.error(f"Error writing to template DOCX: {e}")
        return False

# --- DOCX GENERATOR & HELPER FUNCTIONS ---
def parse_supabase_timestamp(data_list, field_name='timestamp'):
    """Converts ISO 8601 timestamp strings from Supabase into Python datetime objects."""
    for item in data_list:
        if item.get(field_name):
            # Handle potential timezone info and fractional seconds
            iso_string = item[field_name].split('+')[0].split('.')[0]
            try:
                item[field_name] = datetime.fromisoformat(iso_string)
            except ValueError:
                # Fallback for unexpected formats, though less likely
                pass
    return data_list

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def replace_text_in_paragraph(paragraph, replacements):
    """
    Replaces placeholders in a paragraph with values from the replacements dict.
    Handles converting literal '\n' strings into actual newlines for Word.
    """
    for key, val in replacements.items():
        if key in paragraph.text:
            inline = paragraph.runs
            full_text = "".join(run.text for run in inline)
            if key in full_text:
                for run in inline:
                    if key in run.text:
                        # FIX: Convert literal string "\n" (backslash+n) to actual newline character
                        # This ensures lists and multi-line text render correctly in the DOCX
                        replacement_text = str(val).replace('\\n', '\n')
                        run.text = run.text.replace(key, replacement_text)  

def replace_text_in_cell(cell, replacements):
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

def replace_placeholders(doc, replacements):
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, replacements)
    return doc

def flatten_json(data):
    out = {}
    def flatten(x, name=''):
        if isinstance(x, dict):
            for a in x:
                flatten(x[a], name + a)
        elif isinstance(x, list):
            out[name] = "\n".join(map(str, x))
        else:
            out[name] = x
    flatten(data)
    return out

def create_notification(user_id, message):
    try:
        supabase.table('notifications').insert({
            'user_id': user_id,
            'message': message
        }).execute()
    except PostgrestAPIError as e:
        current_app.logger.error(f"Failed to create notification for user {user_id}: {e.message}")

def get_current_user_profile():
    user_id = session.get('user_id')
    if not user_id: return None
    try:
        res = supabase.table('users').select('*').eq('id', user_id).single().execute()
        return res.data
    except PostgrestAPIError as e:
        current_app.logger.error(f"Could not fetch profile for user {user_id}: {e.message}")
        return None

# --- AI GENERATION BACKGROUND TASK (REFACTORED AND IMPROVED) ---
# Update this function to accept plan_id
def start_clp_generation(plan_id, user_id, course_data):
    """A public function to start the background thread."""
    # Pass plan_id to the background task
    thread = Thread(target=generate_clp_background_task, args=(current_app.app_context(), plan_id, user_id, course_data))
    thread.daemon = True
    thread.start()

def generate_clp_background_task(app_context, plan_id, user_id, course_data):
    subject_name = course_data['subject']
    department = course_data['department']
  
    with app_context:
        try:
            current_app.logger.info(f"--- [AI DEBUG] BG Task started for CLP {plan_id}, user {user_id}. ---")
            
            model_instance = GenerativeModel('gemini-2.5-flash')
            clp_data = {}
            
            # --- Step 1: Generate Basic Info and References (Text Generation) ---

            # --- Step 2: Generate Program to Institutional Outcomes Mapping ---
            current_app.logger.info("--- [AI DEBUG] Step 2: Generating Program to Institutional Outcomes Mapping... ---")
            po_io_prompt_raw = get_system_prompt('prompt_po_io', default_text="You are an expert academic planner...") 
            po_io_schema_properties = {f"{po_code}_{io_header}": {"type": "STRING", "enum": ["✔", " "]} for po_code in PROGRAM_OUTCOMES_HEADERS for io_header in INSTITUTIONAL_OUTCOMES_HEADERS}
            step2_config = {"response_mime_type": "application/json", "response_schema": {"type": "OBJECT", "properties": po_io_schema_properties, "required": list(po_io_schema_properties.keys())}}
            resp_2 = model_instance.generate_content(contents=[po_io_prompt_raw], generation_config=step2_config)
            clp_data.update(json.loads(resp_2.text))

            # [STEP 3: CO-PO]
            course_outcomes_string = ", ".join([f"{co['code']}: {co['description']}" for co in COURSE_OUTCOMES])
            program_outcomes_string = ", ".join([f"{po['code']}: {po['description']}" for po in PROGRAM_OUTCOMES])
            co_po_prompt_raw = get_system_prompt('prompt_co_po', default_text="Given the Course Outcomes...")
            co_po_prompt = co_po_prompt_raw.replace('{course_outcomes}', course_outcomes_string).replace('{program_outcomes}', program_outcomes_string)
            co_po_schema_properties = {f"{co_code_obj['code']}_{po_code}": {"type": "STRING", "enum": ["E", "I", " "]} for co_code_obj in COURSE_OUTCOMES for po_code in PROGRAM_OUTCOMES_HEADERS}
            step3_config = {"response_mime_type": "application/json", "response_schema": {"type": "OBJECT", "properties": co_po_schema_properties, "required": list(co_po_schema_properties.keys())}}
            resp_3 = model_instance.generate_content(contents=[co_po_prompt], generation_config=step3_config)
            clp_data.update(json.loads(resp_3.text))

            # [STEP 4: Weekly]
            weekly_prompt_raw = get_system_prompt('prompt_weekly', default_text="Generate the complete 18-week Course Outline...")
            weekly_breakdown_prompt = weekly_prompt_raw.replace('{subject_name}', subject_name)
            weekly_schema_properties = {}
            week_prefixes = [f"W{i}" for i in range(1, 19) if i not in [10, 11, 14, 15, 16, 17]] + ["W1011", "W1415", "W1617"]
            for week_prefix in week_prefixes:
                for suffix in ["_LO", "_TO", "_Method", "_Assesment", "_LR"]:
                    weekly_schema_properties[f"{week_prefix}{suffix}"] = {"type": "STRING"}
            weekly_schema_properties['references'] = {"type": "STRING"}
            step4_config = {"response_mime_type": "application/json", "response_schema": {"type": "OBJECT", "properties": weekly_schema_properties, "required": list(weekly_schema_properties.keys())}}
            resp_4 = model_instance.generate_content(contents=[weekly_breakdown_prompt], generation_config=step4_config)
            clp_data.update(json.loads(resp_4.text))

            # Inject User Data
            clp_data.update(course_data)
            
            # --- SUPABASE PROCESSING ---
            
            # 1. Fetch user profile for template filling
            user_res = supabase.table('users').select('first_name, last_name, title').eq('id', user_id).single().execute()
            current_user = user_res.data
            clp_data['NAME'] = f"{current_user.get('first_name', '')} {current_user.get('last_name', '')}".strip()
            clp_data['TITLE'] = current_user.get('title', '')
            
            # 2. Select Template (Dynamic Logic)
            template_key = None
            try:
                dept_res = supabase.table('departments').select('id').eq('name', department).execute()
                if dept_res.data:
                    dept_id = dept_res.data[0]['id']
                    tmpl_res = supabase.table('templates').select('filename').eq('department_id', dept_id).limit(1).execute()
                    if tmpl_res.data:
                        template_key = tmpl_res.data[0]['filename']
                
                if not template_key:
                    def_res = supabase.table('templates').select('filename').eq('is_default', True).limit(1).execute()
                    if def_res.data:
                        template_key = def_res.data[0]['filename']
                
                if not template_key:
                    template_key = "PBSIT/PBSIT-001-LP-20242.docx" # Fallback

                template_bytes = supabase.storage.from_(STORAGE_BUCKET_NAME).download(template_key)
            except Exception as e:
                # Fallback
                template_key = "PBSIT/PBSIT-001-LP-20242.docx"
                template_bytes = supabase.storage.from_(STORAGE_BUCKET_NAME).download(template_key)

            # 3. Generate DOCX
            doc = Document(io.BytesIO(template_bytes))
            flat_replacements = flatten_json(clp_data)
            doc = replace_placeholders(doc, flat_replacements)
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            # 4. Upload File
            # Use plan_id in filename to ensure uniqueness
            docx_filename = f"{subject_name.replace(' ', '_')}_Generated_{plan_id}.docx"
            file_path_in_bucket = f"{user_id}/{docx_filename}"
            
            supabase.storage.from_(STORAGE_BUCKET_NAME).upload(
                path=file_path_in_bucket,
                file=file_stream.read(),
                file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            )
            
            # 5. UPDATE the existing database record (don't insert new)
            current_app.logger.info(f"--- [AI DEBUG] BG Task: Finalizing CLP {plan_id}... ---")
            
            final_subject = clp_data.get('descriptive_title', subject_name)
            
            supabase.table('course_learning_plans').update({
                'subject': final_subject, # Update subject if AI refined it
                'filename': file_path_in_bucket,
                'upload_type': 'file_upload', # Switch to file type so it opens in ONLYOFFICE
                'content': json.dumps(clp_data), # Keep JSON for backup/viewing
                'status': 'draft' # Mark as done (Draft allows editing)
            }).eq('id', plan_id).execute()
            
            # 6. Notify
            create_notification(user_id, f'Your AI-generated CLP for "{final_subject}" is ready!')
            
        except Exception as e:
            current_app.logger.error(f"--- [AI DEBUG] BG Task FAILED: {e} ---")
            import traceback
            current_app.logger.error(traceback.format_exc())
            
            # Mark as failed in DB so UI stops loading
            try:
                supabase.table('course_learning_plans').update({
                    'status': 'failed', # You might need to ensure UI handles 'failed' or just 'draft' with error note
                    'content': f"Generation failed: {str(e)}"
                }).eq('id', plan_id).execute()
            except:
                pass
                
            create_notification(user_id, f'CLP generation failed: {e}')
            
        except Exception as e:
            current_app.logger.error(f"--- [AI DEBUG] BG Task FAILED: {e} ---")
            import traceback
            current_app.logger.error(traceback.format_exc())
            
            # Mark as failed in DB so UI stops loading
            try:
                supabase.table('course_learning_plans').update({
                    'status': 'failed', # You might need to ensure UI handles 'failed' or just 'draft' with error note
                    'content': f"Generation failed: {str(e)}"
                }).eq('id', plan_id).execute()
            except:
                pass
                
            create_notification(user_id, f'CLP generation failed: {e}')
            current_app.logger.info("--- [AI DEBUG] BG Task: Process completed successfully. ---")
            
        except FileNotFoundError as e:
            current_app.logger.error(f"--- [AI DEBUG] BG Task FAILED: {e} ---")
            create_notification(user_id, f'CLP generation for "{subject_name}" failed: {e}. Please contact an administrator.')
        except ValueError as e:
            current_app.logger.error(f"--- [AI DEBUG] BG Task FAILED: {e} ---")
            create_notification(user_id, f'CLP generation failed due to AI content policy. Please try a different subject.')
        except RuntimeError as e:
            current_app.logger.error(f"--- [AI DEBUG] BG Task FAILED: {e} ---")
            create_notification(user_id, f'CLP generation failed due to an AI service error. Please try again later.')
        except Exception as e:
            current_app.logger.error(f"--- [AI DEBUG] BG Task FAILED for user {user_id}: {e} ---")
            # Log the full exception for detailed debugging
            current_app.logger.error("Error on line {}: {}".format(sys.exc_info()[-1].tb_lineno, e))
            create_notification(user_id, f'CLP generation for "{subject_name}" failed unexpectedly. An administrator has been notified.')

def get_system_prompt(key, default_text=""):
    try:
        # Fetch from DB
        res = supabase.table('system_settings').select('value').eq('key', key).single().execute()
        if res.data:
            return res.data['value']
    except Exception as e:
        current_app.logger.error(f"Error fetching prompt '{key}': {e}")
    # Fallback to default if DB fails
    return default_text